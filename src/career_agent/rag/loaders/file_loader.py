"""Multi-format file loader — PDF, DOCX, TXT, MD → ProfileDocument."""

from __future__ import annotations

import hashlib
from pathlib import Path
from typing import Any

from career_agent.rag.schemas import ProfileDocument


class FileLoader:
    """Load PDF, DOCX, TXT, MD files into ProfileDocument objects.

    Usage::

        loader = FileLoader()
        doc = loader.load("resume.pdf")
        pipeline.build_index_from_docs([doc])
    """

    def load(self, file_path: str | Path) -> ProfileDocument:
        path = Path(file_path).resolve()
        if not path.is_file():
            raise FileNotFoundError(f"File not found: {path}")

        suffix = path.suffix.lower()
        if suffix == ".pdf":
            content = self._read_pdf(path)
        elif suffix in (".docx", ".doc"):
            content = self._read_docx(path)
        elif suffix in (".txt", ".md", ".markdown"):
            content = path.read_text(encoding="utf-8")
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

        doc_id = hashlib.sha256(str(path).encode()).hexdigest()[:16]
        return ProfileDocument(
            document_id=doc_id,
            source_path=str(path),
            title=path.stem,
            content=content,
            item_type=self._infer_type(path),
            metadata={
                "filename": path.name,
                "file_stem": path.stem,
                "source_path": str(path),
                "loader": "file_loader",
                "file_type": suffix,
            },
        )

    def load_directory(self, dir_path: str | Path) -> list[ProfileDocument]:
        """Load all supported files from a directory."""
        path = Path(dir_path).resolve()
        if not path.is_dir():
            return []
        docs = []
        supported = {".pdf", ".docx", ".doc", ".txt", ".md", ".markdown"}
        for f in sorted(path.iterdir()):
            if f.is_file() and f.suffix.lower() in supported:
                try:
                    docs.append(self.load(f))
                except Exception:
                    pass  # skip corrupted files
        return docs

    @staticmethod
    def _read_pdf(path: Path) -> str:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
        return "\n\n".join(parts)

    @staticmethod
    def _read_docx(path: Path) -> str:
        from docx import Document
        doc = Document(str(path))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        return "\n\n".join(parts)

    @staticmethod
    def _infer_type(path: Path) -> str:
        stem = path.stem.lower()
        type_map = {
            "resume": "resume", "简历": "resume", "cv": "resume",
            "projects": "project", "项目": "project",
            "skills": "skill", "技能": "skill",
            "internship": "internship", "实习": "internship",
        }
        for key, val in type_map.items():
            if key in stem:
                return val
        return "document"
