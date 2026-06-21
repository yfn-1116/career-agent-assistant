"""DOCX resume exporter — generates .docx from match results."""

from __future__ import annotations
from pathlib import Path
from typing import Any


def export_docx(
    output_path: str,
    profile_info: dict[str, Any] | None = None,
    match_result: dict[str, Any] | None = None,
    generated_bullets: list[str] | None = None,
    template: str = "agent_role",
) -> str:
    """Export a .docx resume file.

    Returns the absolute path to the generated file.
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    info = profile_info or {}
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    # Title
    title = doc.add_heading(info.get("name", "姓名"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parts = []
    for key, label in [("school", "学校"), ("major", "专业"), ("email", "邮箱"), ("phone", "电话"), ("github", "GitHub")]:
        if info.get(key):
            parts.append(f"{info[key]}")
    contact.add_run(" | ".join(parts)).font.size = Pt(10)

    # Education
    if info.get("school"):
        doc.add_heading("教育背景", level=1)
        doc.add_paragraph(f"{info.get('school', '')} | {info.get('major', '')} | {info.get('degree', '')}")

    # Skills
    if match_result:
        matched = match_result.get("matched_skills", [])
        if matched:
            doc.add_heading("专业技能", level=1)
            doc.add_paragraph(", ".join(matched[:12]))

    # Projects
    bullets = generated_bullets or []
    if bullets:
        doc.add_heading("项目经历", level=1)
        for i, bullet in enumerate(bullets[:5], 1):
            text = bullet.replace("# ", "").replace("\n", " ")[:250]
            if text.startswith("- "):
                text = text[2:]
            p = doc.add_paragraph(f"{text}", style="List Bullet")

    # Self evaluation
    if match_result:
        strengths = match_result.get("matched_skills", [])
        if strengths:
            doc.add_heading("自我评价", level=1)
            doc.add_paragraph(f"具备 {', '.join(strengths[:5])} 等方面的项目经验和实践能力。")

    # Save
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return str(path.resolve())
